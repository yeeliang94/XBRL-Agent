"""Phase 4 / Step 7: the Download-as-Word endpoint.

Note: Docling has no .docx exporter (HTML/MD/JSON only), so the converted HTML
is rendered to Word via htmldocx. This pins that the download produces a valid
.docx whose table carries the key figures, and that the error states are clean.
"""
from __future__ import annotations

import importlib
import io
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from db.schema import init_db
from db import repository as repo

_HTML = (
    "<html><body><h3>Statement of financial position</h3>"
    "<table><tr><th>Item</th><th>2021</th></tr>"
    "<tr><td>Receivables</td><td>391,675</td></tr>"
    "<tr><td>Total assets</td><td>3,141,738</td></tr></table></body></html>"
)


@pytest.fixture
def client(tmp_path: Path, monkeypatch):
    db = tmp_path / "xbrl.db"
    monkeypatch.setenv("XBRL_OUTPUT_DIR", str(tmp_path))
    monkeypatch.setenv("GOOGLE_API_KEY", "test-key-12345")
    monkeypatch.setenv("LLM_PROXY_URL", "")
    import server as srv
    importlib.reload(srv)
    srv.AUDIT_DB_PATH = db
    init_db(db)
    return TestClient(srv.app), db, tmp_path


def _seed_done_job(db, tmp_path) -> int:
    html_path = tmp_path / "doc_conversions" / "job1" / "result.html"
    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(_HTML, encoding="utf-8")
    with repo.db_session(db) as conn:
        jid = repo.create_doc_conversion(
            conn, source_pdf_path="/tmp/FINCO.pdf", original_filename="FINCO.pdf"
        )
        repo.mark_doc_conversion_finished(
            conn, jid, status="done", result_html_path=str(html_path)
        )
    return jid


def test_download_docx_has_table_with_figures(client):
    from docx import Document

    c, db, tmp_path = client
    jid = _seed_done_job(db, tmp_path)

    resp = c.get(f"/api/doc-convert/{jid}/download/docx")
    assert resp.status_code == 200
    assert "wordprocessingml.document" in resp.headers["content-type"]
    assert "FINCO-readable.docx" in resp.headers["content-disposition"]

    doc = Document(io.BytesIO(resp.content))
    assert len(doc.tables) >= 1
    cells = [c.text for row in doc.tables[0].rows for c in row.cells]
    assert any("3,141,738" in cell for cell in cells)


def test_download_docx_409_before_done(client):
    c, db, tmp_path = client
    with repo.db_session(db) as conn:
        jid = repo.create_doc_conversion(
            conn, source_pdf_path="/tmp/a.pdf", original_filename="a.pdf"
        )
    assert c.get(f"/api/doc-convert/{jid}/download/docx").status_code == 409


def test_download_docx_404_unknown(client):
    c, _, _ = client
    assert c.get("/api/doc-convert/9999/download/docx").status_code == 404


def test_download_docx_sanitizes_unsafe_filename(client, tmp_path):
    """A non-ASCII / quote-bearing filename must not break the header or 500.

    Regression for the code-review finding: the Content-Disposition is built
    from the user-supplied filename, so it must be sanitized (ASCII filename +
    RFC 5987 filename*).
    """
    c, db, tp = client
    html_path = tp / "doc_conversions" / "job2" / "result.html"
    html_path.parent.mkdir(parents=True, exist_ok=True)
    html_path.write_text(_HTML, encoding="utf-8")
    with repo.db_session(db) as conn:
        jid = repo.create_doc_conversion(
            conn, source_pdf_path="/tmp/x.pdf", original_filename='财务报表"evil.pdf'
        )
        repo.mark_doc_conversion_finished(
            conn, jid, status="done", result_html_path=str(html_path)
        )

    resp = c.get(f"/api/doc-convert/{jid}/download/docx")
    assert resp.status_code == 200  # NOT a 500 from a latin-1 header encode error
    cd = resp.headers["content-disposition"]
    # ASCII filename carries no raw quote/unicode; the unicode rides on filename*.
    assert 'filename="' in cd
    assert '财' not in cd.split("filename*=")[0]  # no unicode in the ascii part
    assert "filename*=UTF-8''" in cd
